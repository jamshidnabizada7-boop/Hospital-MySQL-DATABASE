import docx

doc = docx.Document(r"d:\Hospital MYSQL Databse\Hospital_Management_System_Report.docx")

styles = set(p.style.name for p in doc.paragraphs)
print("Unique paragraph styles in DOCX:", styles)

print("\nFirst 20 paragraphs in DOCX:")
for i, p in enumerate(doc.paragraphs[:20]):
    if p.text.strip():
        print(f"P{i+1} [Style: {p.style.name}]: {p.text[:80]}")
