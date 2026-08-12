"""
Empirical Verification Script for Hospital Management System Report Files
Created by Challenger Report 2 agent.
"""

import os
import sys
import re
import xml.etree.ElementTree as ET
import zipfile

MD_PATH = r"d:\Hospital MYSQL Databse\project_report.md"
DOCX_PATH = r"d:\Hospital MYSQL Databse\Hospital_Management_System_Report.docx"

REQUIRED_SECTIONS = [
    ("Abstract", ["abstract", "executive summary"]),
    ("System Architecture", ["system architecture", "technologies used", "architecture topology"]),
    ("Database Schema", ["database schema", "schema & design", "er structure", "tables", "relationships"]),
    ("Access Control", ["access control", "security features", "role access control"]),
    ("Frontend UI Flow", ["frontend ui flow", "ui flow", "component architecture"]),
    ("Future Enhancements", ["future enhancements", "conclusion"])
]

def run_empirical_verification():
    print("======================================================================")
    print("EMPIRICAL VERIFICATION REPORT -- CHALLENGER 2")
    print("======================================================================\n")

    # 1. MARKDOWN FILE VERIFICATION
    print("--- 1. MARKDOWN REPORT (project_report.md) ---")
    if not os.path.exists(MD_PATH):
        print(f"FAIL: Markdown file missing at {MD_PATH}")
        sys.exit(1)
    
    md_size = os.path.getsize(MD_PATH)
    with open(MD_PATH, 'r', encoding='utf-8', errors='ignore') as f:
        md_text = f.read()

    md_chars = len(md_text)
    md_lines = md_text.splitlines()
    md_non_empty_lines = [l for l in md_lines if l.strip()]
    md_paras = [p.strip() for p in re.split(r'\n\s*\n', md_text) if p.strip()]
    
    # Headings
    headings = re.findall(r'^(#+)\s+(.+)$', md_text, re.MULTILINE)
    h_counts = {1: 0, 2: 0, 3: 0, 4: 0}
    for h in headings:
        lvl = min(len(h[0]), 4)
        h_counts[lvl] += 1
        
    # Tables & SQL
    md_table_blocks = 0
    in_table = False
    for line in md_lines:
        if '|' in line:
            if not in_table:
                md_table_blocks += 1
                in_table = True
        else:
            in_table = False
            
    sql_blocks = re.findall(r'```sql(.*?)```', md_text, re.DOTALL | re.IGNORECASE)
    all_code_blocks = re.findall(r'```(.*?)```', md_text, re.DOTALL)

    print(f"[*] File Size: {md_size:,} bytes")
    print(f"[*] Character Count: {md_chars:,} (Threshold: >50,000) -> {'PASS' if md_chars > 50000 else 'FAIL'}")
    print(f"[*] Non-empty Lines: {len(md_non_empty_lines):,} (Line Paras)")
    print(f"[*] Blank-separated Paragraph Blocks: {len(md_paras):,}")
    print(f"[*] Total Headings: {len(headings)} (H1:{h_counts[1]}, H2:{h_counts[2]}, H3:{h_counts[3]}, H4+:{h_counts[4]})")
    print(f"[*] Table Blocks: {md_table_blocks} (Threshold: >0) -> {'PASS' if md_table_blocks > 0 else 'FAIL'}")
    print(f"[*] SQL Code Blocks: {len(sql_blocks)} total (All Code Blocks: {len(all_code_blocks)}) -> {'PASS' if len(sql_blocks) > 0 else 'FAIL'}")

    # Check Required Sections in MD
    md_missing_sections = []
    print("\nSection Presence in Markdown:")
    for sec_key, keywords in REQUIRED_SECTIONS:
        found_heading = None
        for h_level, h_title in headings:
            if any(kw in h_title.lower() for kw in keywords):
                found_heading = h_title
                break
        if found_heading:
            print(f"  - {sec_key}: PASS (Matched heading: '{found_heading}')")
        else:
            print(f"  - {sec_key}: FAIL (Not found)")
            md_missing_sections.append(sec_key)

    # 2. DOCX FILE VERIFICATION
    print("\n--- 2. DOCX REPORT (Hospital_Management_System_Report.docx) ---")
    if not os.path.exists(DOCX_PATH):
        print(f"FAIL: DOCX file missing at {DOCX_PATH}")
        sys.exit(1)

    docx_size = os.path.getsize(DOCX_PATH)
    import docx
    doc = docx.Document(DOCX_PATH)

    docx_paras = doc.paragraphs
    docx_non_empty_paras = [p for p in docx_paras if p.text.strip()]
    docx_tables = doc.tables

    full_text_list = [p.text for p in docx_paras]
    for t in docx_tables:
        for row in t.rows:
            for cell in row.cells:
                full_text_list.append(cell.text)
    docx_full_text = "\n".join(full_text_list)
    docx_chars = len(docx_full_text)

    sql_kw_matches = [kw for kw in ["CREATE TABLE", "FOREIGN KEY", "PRIMARY KEY", "SELECT", "INSERT INTO", "JOIN"] if kw in docx_full_text.upper()]

    print(f"[*] File Size: {docx_size:,} bytes")
    print(f"[*] Character Count: {docx_chars:,} (Threshold: >50,000) -> {'PASS' if docx_chars > 50000 else 'FAIL'}")
    print(f"[*] Non-empty Paragraph Count: {len(docx_non_empty_paras):,} (Threshold: >1,000) -> {'PASS' if len(docx_non_empty_paras) > 1000 else 'FAIL'}")
    print(f"[*] Table Count: {len(docx_tables)} (Threshold: >0) -> {'PASS' if len(docx_tables) > 0 else 'FAIL'}")
    print(f"[*] SQL Keywords Present: {sql_kw_matches} -> {'PASS' if sql_kw_matches else 'FAIL'}")

    # Check Required Sections in DOCX
    docx_missing_sections = []
    print("\nSection Presence in DOCX:")
    for sec_key, keywords in REQUIRED_SECTIONS:
        found = any(kw in docx_full_text.lower() for kw in keywords)
        if found:
            print(f"  - {sec_key}: PASS")
        else:
            print(f"  - {sec_key}: FAIL")
            docx_missing_sections.append(sec_key)

    # 3. VERDICT
    print("\n======================================================================")
    all_passed = (
        md_chars > 50000 and
        len(sql_blocks) > 0 and
        md_table_blocks > 0 and
        len(md_missing_sections) == 0 and
        docx_chars > 50000 and
        len(docx_non_empty_paras) > 1000 and
        len(docx_tables) > 0 and
        len(sql_kw_matches) > 0 and
        len(docx_missing_sections) == 0
    )

    verdict = "APPROVE" if all_passed else "REQUEST_CHANGES"
    print(f"VERDICT: {verdict}")
    print("======================================================================")
    return verdict

if __name__ == "__main__":
    run_empirical_verification()
