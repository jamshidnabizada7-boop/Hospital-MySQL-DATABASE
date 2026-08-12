import docx
import os
import sys

def verify():
    docx_path = r"d:\Hospital MYSQL Databse\Hospital_Management_System_Report.docx"
    md_path = r"d:\Hospital MYSQL Databse\project_report.md"

    print(f"Verifying DOCX: {docx_path}")
    print(f"Verifying MD: {md_path}")

    assert os.path.exists(docx_path), "DOCX file does not exist!"
    assert os.path.exists(md_path), "MD file does not exist!"

    doc = docx.Document(docx_path)
    with open(md_path, 'r', encoding='utf-8') as f:
        md_text = f.read()

    print(f"DOCX Paragraphs: {len(doc.paragraphs)}")
    print(f"DOCX Tables: {len(doc.tables)}")

    # Check 6 required sections in MD and DOCX
    sections = [
        "Abstract",
        "System Architecture",
        "Database Schema",
        "Access Control",
        "Frontend UI Flow",
        "Future Enhancements"
    ]

    docx_full_text = "\n".join([p.text for p in doc.paragraphs])
    
    print("\n--- Section Presence Check ---")
    for sec in sections:
        in_md = sec.lower() in md_text.lower()
        in_docx = sec.lower() in docx_full_text.lower()
        print(f"Section '{sec}': MD={in_md}, DOCX={in_docx}")

    # Check table details
    print("\n--- DOCX Table Inspection ---")
    for idx, tbl in enumerate(doc.tables):
        print(f"Table {idx+1}: {len(tbl.rows)} rows, {len(tbl.columns)} cols")
        hdr = [cell.text.strip().replace('\n', ' ') for cell in tbl.rows[0].cells]
        print(f"  Header: {hdr[:4]}")

    # Check database entities in MD and DOCX
    print("\n--- Database Content Check ---")
    db_items = [
        "23 normalized tables", "23 tables",
        "Role", "App_User", "Department", "Doctor", "Employee", "Patient",
        "Doctor_Schedule", "Appointment_Slot", "Appointment", "Medical_Record",
        "Prescription", "Prescription_Item", "Medicine_Category", "Medicine",
        "Pharmacy", "Inventory", "Lab_Test", "Lab_Order", "Lab_Result",
        "Bill", "Payment", "Audit_Log"
    ]
    for item in db_items:
        found_md = item in md_text
        found_docx = item in docx_full_text
        if not (found_md and found_docx):
            print(f"  DB Item '{item}': MD={found_md}, DOCX={found_docx}")

    print("\n--- Code Block Sample Check in DOCX ---")
    code_samples = ["CREATE TABLE", "SELECT ... FOR UPDATE", "DELIMITER $$", "jwt.verify", "pushState"]
    for cs in code_samples:
        print(f"  Code snippet '{cs}': in_MD={cs in md_text}, in_DOCX={cs in docx_full_text}")

if __name__ == "__main__":
    verify()
