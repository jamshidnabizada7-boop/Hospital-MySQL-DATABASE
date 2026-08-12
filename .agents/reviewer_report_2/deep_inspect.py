import docx
import os
import sys

# Force UTF-8 encoding for stdout
sys.stdout.reconfigure(encoding='utf-8')

def deep_inspect_docx():
    docx_path = r"d:\Hospital MYSQL Databse\Hospital_Management_System_Report.docx"
    doc = docx.Document(docx_path)

    print("--- DEEP DOCX INSPECTION ---")
    print(f"File Size: {os.path.getsize(docx_path):,} bytes")
    print(f"Total Paragraphs: {len(doc.paragraphs)}")
    print(f"Total Tables: {len(doc.tables)}")

    # Inspect font names and sizes in runs
    fonts = set()
    sizes = set()
    code_runs = 0
    for p in doc.paragraphs:
        for r in p.runs:
            if r.font.name:
                fonts.add(r.font.name)
            if r.font.size:
                sizes.add(r.font.size.pt)
            if r.font.name == 'Consolas':
                code_runs += 1

    print(f"Fonts detected in document runs: {sorted(list(fonts))}")
    print(f"Font sizes (pt) detected: {sorted(list(sizes))}")
    print(f"Code runs count (Consolas font): {code_runs}")

    # Inspect tables
    for i, t in enumerate(doc.tables):
        print(f"\nTable {i+1}:")
        print(f"  Rows: {len(t.rows)}, Cols: {len(t.columns)}")
        print(f"  Header text: {[cell.text.strip() for cell in t.rows[0].cells]}")

    # 6 Main Section Headings Check
    main_sections = [
        "1. Abstract & Executive Summary",
        "2. System Architecture & Technologies Used",
        "3. Database Schema & Design",
        "4. Access Control & Security Features",
        "5. Frontend UI Flow & Component Architecture",
        "6. Future Enhancements & Conclusion"
    ]

    doc_text = "\n".join([p.text for p in doc.paragraphs])
    print("\nMain 6 Sections Check:")
    for sec in main_sections:
        present = sec in doc_text
        print(f"  [{'PASS' if present else 'FAIL'}] {sec}")

if __name__ == "__main__":
    deep_inspect_docx()
