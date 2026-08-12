import os
import sys
import docx

def verify_docx(filepath):
    print("==================================================")
    print(f"Verifying document: {filepath}")
    print("==================================================")

    # 1. Existence check
    if not os.path.exists(filepath):
        print(f"FAIL: File does not exist: {filepath}")
        return False
    print("[PASS] File exists.")

    # 2. Non-empty check
    filesize = os.path.getsize(filepath)
    print(f"File size: {filesize} bytes")
    if filesize <= 0:
        print("FAIL: File is empty (0 bytes).")
        return False
    print("[PASS] File is non-empty (> 0 bytes).")

    # 3. Readability & validity check
    try:
        doc = docx.Document(filepath)
    except Exception as e:
        print(f"FAIL: Unable to open document with python-docx: {e}")
        return False
    print("[PASS] File is a valid, readable Microsoft Word (.docx) document.")

    # Extract all text from paragraphs and table cells
    all_text_list = []
    for p in doc.paragraphs:
        all_text_list.append(p.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    all_text_list.append(p.text)

    full_document_text = "\n".join(all_text_list)

    print(f"Total Paragraphs: {len(doc.paragraphs)}")
    print(f"Total Tables: {len(doc.tables)}")
    print(f"Total Text Character Length: {len(full_document_text)}")

    # 4. Section Presence Checks
    required_sections = [
        ("Abstract / Executive Summary", ["Abstract", "Executive Summary"]),
        ("System Architecture & Technologies Used", ["System Architecture", "Technologies Used"]),
        ("Database Schema", ["Database Schema", "ER", "Tables", "Relationships", "Constraints", "Complex Queries"]),
        ("Access Control & Security Features", ["Access Control", "Security Features"]),
        ("Frontend UI Flow", ["Frontend UI Flow", "Component Architecture"]),
        ("Future Enhancements & Conclusion", ["Future Enhancements", "Conclusion"])
    ]

    missing_sections = []

    for sec_name, keywords in required_sections:
        found = False
        for kw in keywords:
            if kw.lower() in full_document_text.lower():
                found = True
                break
        if found:
            print(f"[PASS] Found required section concept: '{sec_name}'")
        else:
            print(f"[FAIL] Missing required section concept: '{sec_name}'")
            missing_sections.append(sec_name)

    if missing_sections:
        print("\nVerification FAILED. Missing sections:", missing_sections)
        return False

    print("\n==================================================")
    print("ALL VERIFICATION CHECKS PASSED SUCCESSFULLY!")
    print("==================================================")
    return True

if __name__ == '__main__':
    target_file = r"d:\Hospital MYSQL Databse\Hospital_Management_System_Report.docx"
    success = verify_docx(target_file)
    if not success:
        sys.exit(1)
