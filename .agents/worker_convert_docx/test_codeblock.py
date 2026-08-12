import docx
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

doc = docx.Document()

diagram = """+-----------------------------------------------------------------------------------+
|                                PRESENTATION TIER                                  |
|  Single Page Application (SPA) - HTML5 / CSS3 / Vanilla JS (ES6+)                 |
|  - Client Routing (pushState/popstate)         - 9 Interactive Domain Views       |
+-----------------------------------------------------------------------------------+"""

lines = diagram.split('\n')
for i, line in enumerate(lines):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1.05
    p.paragraph_format.left_indent = Inches(0.25)
    
    if i < len(lines) - 1:
        p.paragraph_format.keep_with_next = True
        
    pPr = p._p.get_or_add_pPr()
    shd = parse_xml(r'<w:shd {} w:fill="F1F5F9"/>'.format(nsdecls('w')))
    pPr.append(shd)
    
    if i == 0:
        pBdr = parse_xml(r'<w:pBdr {}><w:top w:val="single" w:sz="4" w:space="4" w:color="CBD5E1"/><w:left w:val="single" w:sz="12" w:space="8" w:color="2563EB"/></w:pBdr>'.format(nsdecls('w')))
        pPr.append(pBdr)
    elif i == len(lines) - 1:
        pBdr = parse_xml(r'<w:pBdr {}><w:bottom w:val="single" w:sz="4" w:space="4" w:color="CBD5E1"/><w:left w:val="single" w:sz="12" w:space="8" w:color="2563EB"/></w:pBdr>'.format(nsdecls('w')))
        pPr.append(pBdr)
    else:
        pBdr = parse_xml(r'<w:pBdr {}><w:left w:val="single" w:sz="12" w:space="8" w:color="2563EB"/></w:pBdr>'.format(nsdecls('w')))
        pPr.append(pBdr)
        
    run = p.add_run(line)
    run.font.name = 'Consolas'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(15, 23, 42)

doc.save("test_codeblock.docx")
print("Saved test_codeblock.docx!")
