import docx
from docx.shared import Inches, Pt, RGBColor
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

doc = docx.Document()
p = doc.add_paragraph("Test paragraph")
# Test shading
shading_elm = parse_xml(r'<w:shd {} w:fill="F1F5F9"/>'.format(nsdecls('w')))
p._p.get_or_add_pPr().append(shading_elm)

table = doc.add_table(rows=2, cols=2)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = "Header 1"
hdr_cells[1].text = "Header 2"

# Shading header cell
for cell in hdr_cells:
    shd = parse_xml(r'<w:shd {} w:fill="1E3A8A"/>'.format(nsdecls('w')))
    cell._tc.get_or_add_tcPr().append(shd)

doc.save("test.docx")
print("Saved test.docx successfully!")
