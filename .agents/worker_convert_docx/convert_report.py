import os
import re
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

def parse_inline_markdown(text):
    """
    Parses inline markdown tokens (**bold**, *italic*, `code`, ***bold-italic***)
    and returns a list of dictionaries: [{'text': str, 'bold': bool, 'italic': bool, 'code': bool}]
    """
    if not text:
        return []

    pattern = re.compile(
        r'`([^`]+)`|'
        r'\*\*\*([^*]+)\*\*\*|'
        r'\*\*([^*]+)\*\*|'
        r'__([^_]+)__|'
        r'\*([^*]+)\*|'
        r'_([^_]+)_'
    )

    runs = []
    last_idx = 0

    for match in pattern.finditer(text):
        start, end = match.span()
        if start > last_idx:
            runs.append({
                'text': text[last_idx:start],
                'bold': False,
                'italic': False,
                'code': False
            })

        code_text = match.group(1)
        bold_italic_text = match.group(2)
        bold_text1 = match.group(3)
        bold_text2 = match.group(4)
        italic_text1 = match.group(5)
        italic_text2 = match.group(6)

        if code_text is not None:
            runs.append({'text': code_text, 'bold': False, 'italic': False, 'code': True})
        elif bold_italic_text is not None:
            runs.append({'text': bold_italic_text, 'bold': True, 'italic': True, 'code': False})
        elif bold_text1 is not None or bold_text2 is not None:
            t = bold_text1 if bold_text1 is not None else bold_text2
            runs.append({'text': t, 'bold': True, 'italic': False, 'code': False})
        elif italic_text1 is not None or italic_text2 is not None:
            t = italic_text1 if italic_text1 is not None else italic_text2
            runs.append({'text': t, 'bold': False, 'italic': True, 'code': False})

        last_idx = end

    if last_idx < len(text):
        runs.append({
            'text': text[last_idx:],
            'bold': False,
            'italic': False,
            'code': False
        })

    return runs


def parse_markdown_blocks(lines):
    """
    Parses lines of markdown file into structured block tokens:
    - {'type': 'heading', 'level': int, 'text': str}
    - {'type': 'code_block', 'lang': str, 'lines': [str]}
    - {'type': 'table', 'headers': [str], 'alignments': [str], 'rows': [[str]]}
    - {'type': 'list_item', 'ordered': bool, 'num': str, 'indent': int, 'text': str}
    - {'type': 'hr'}
    - {'type': 'paragraph', 'text': str}
    """
    blocks = []
    i = 0
    num_lines = len(lines)

    while i < num_lines:
        line = lines[i]
        rstrip_line = line.rstrip()
        strip_line = line.strip()

        # 1. Blank line
        if not strip_line:
            i += 1
            continue

        # 2. Fenced code block
        if strip_line.startswith("```"):
            lang = strip_line[3:].strip()
            code_lines = []
            i += 1
            while i < num_lines:
                if lines[i].strip().startswith("```"):
                    i += 1
                    break
                code_lines.append(lines[i].rstrip("\r\n"))
                i += 1
            blocks.append({
                'type': 'code_block',
                'lang': lang,
                'lines': code_lines
            })
            continue

        # 3. Horizontal rule
        if re.match(r'^(---|\*\*\*|___)\s*$', strip_line):
            blocks.append({'type': 'hr'})
            i += 1
            continue

        # 4. Headings
        heading_match = re.match(r'^(#{1,6})\s+(.*)$', strip_line)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            blocks.append({
                'type': 'heading',
                'level': level,
                'text': text
            })
            i += 1
            continue

        # 5. Table
        if '|' in strip_line and i + 1 < num_lines and '|' in lines[i + 1] and ('---' in lines[i + 1] or ':-' in lines[i + 1] or '-:' in lines[i + 1]):
            # Header line
            header_cells = [c.strip() for c in strip_line.strip('|').split('|')]
            # Alignments line
            align_line = lines[i + 1].strip().strip('|')
            align_raw = [c.strip() for c in align_line.split('|')]
            alignments = []
            for a in align_raw:
                if a.startswith(':') and a.endswith(':'):
                    alignments.append('center')
                elif a.endswith(':'):
                    alignments.append('right')
                else:
                    alignments.append('left')

            i += 2
            data_rows = []
            while i < num_lines and '|' in lines[i] and lines[i].strip():
                row_cells = [c.strip() for c in lines[i].strip().strip('|').split('|')]
                # Ensure correct column count matching header
                if len(row_cells) < len(header_cells):
                    row_cells.extend([''] * (len(header_cells) - len(row_cells)))
                elif len(row_cells) > len(header_cells):
                    row_cells = row_cells[:len(header_cells)]
                data_rows.append(row_cells)
                i += 1

            blocks.append({
                'type': 'table',
                'headers': header_cells,
                'alignments': alignments,
                'rows': data_rows
            })
            continue

        # 6. List items (Unordered & Ordered)
        ul_match = re.match(r'^(\s*)([-*+])\s+(.*)$', rstrip_line)
        ol_match = re.match(r'^(\s*)(\d+)\.\s+(.*)$', rstrip_line)

        if ul_match:
            indent_spaces = len(ul_match.group(1))
            indent_level = indent_spaces // 2
            text = ul_match.group(3).strip()
            blocks.append({
                'type': 'list_item',
                'ordered': False,
                'num': '',
                'indent': indent_level,
                'text': text
            })
            i += 1
            continue
        elif ol_match:
            indent_spaces = len(ol_match.group(1))
            indent_level = indent_spaces // 2
            num_str = ol_match.group(2)
            text = ol_match.group(3).strip()
            blocks.append({
                'type': 'list_item',
                'ordered': True,
                'num': num_str,
                'indent': indent_level,
                'text': text
            })
            i += 1
            continue

        # 7. Paragraph (accumulate consecutive text lines until blank or block element)
        para_lines = []
        while i < num_lines:
            curr_line = lines[i].strip()
            if not curr_line:
                break
            if curr_line.startswith("```") or curr_line.startswith("#") or re.match(r'^(---|\*\*\*|___)\s*$', curr_line):
                break
            if '|' in curr_line and i + 1 < num_lines and '|' in lines[i + 1] and ('---' in lines[i + 1] or ':-' in lines[i + 1] or '-:' in lines[i + 1]):
                break
            if re.match(r'^\s*[-*+]\s+', lines[i]) or re.match(r'^\s*\d+\.\s+', lines[i]):
                break
            para_lines.append(curr_line)
            i += 1

        if para_lines:
            blocks.append({
                'type': 'paragraph',
                'text': ' '.join(para_lines)
            })

    return blocks


def add_runs_to_paragraph(p, runs, default_font='Calibri', default_size=11, default_color=RGBColor(30, 41, 59)):
    """
    Appends inline runs to a docx paragraph with appropriate fonts and formatting.
    """
    for r in runs:
        run = p.add_run(r['text'])
        run.bold = r['bold']
        run.italic = r['italic']

        if r['code']:
            run.font.name = 'Consolas'
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(37, 99, 235)  # Royal blue for code span
            # Light background for inline code via XML
            rPr = run._r.get_or_add_rPr()
            shd = parse_xml(r'<w:shd {} w:fill="F1F5F9"/>'.format(nsdecls('w')))
            rPr.append(shd)
        else:
            run.font.name = default_font
            run.font.size = Pt(default_size)
            run.font.color.rgb = default_color


def build_docx(blocks, output_filepath):
    """
    Constructs a professionally styled .docx document from parsed block tokens.
    """
    doc = docx.Document()

    # Set 1-inch Page Margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

        # Header
        header = section.header
        hp = header.paragraphs[0]
        hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        hrun = hp.add_run("Hospital Management System | Technical Project Report")
        hrun.font.name = 'Calibri'
        hrun.font.size = Pt(8.5)
        hrun.font.color.rgb = RGBColor(100, 116, 139)

        # Footer
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        frun = fp.add_run("Database Management Systems — Enterprise HMS Solution")
        frun.font.name = 'Calibri'
        frun.font.size = Pt(8.5)
        frun.font.color.rgb = RGBColor(100, 116, 139)

    is_first_h1 = True

    for block in blocks:
        btype = block['type']

        if btype == 'heading':
            level = block['level']
            text = block['text']

            p = doc.add_paragraph()
            p.paragraph_format.keep_with_next = True

            if level == 1:
                if is_first_h1:
                    # Document Title
                    is_first_h1 = False
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(12)
                    runs = parse_inline_markdown(text)
                    add_runs_to_paragraph(p, runs, default_font='Segoe UI', default_size=22, default_color=RGBColor(30, 58, 138))
                else:
                    # Section Heading (H1)
                    p.paragraph_format.space_before = Pt(16)
                    p.paragraph_format.space_after = Pt(6)
                    runs = parse_inline_markdown(text)
                    add_runs_to_paragraph(p, runs, default_font='Segoe UI', default_size=16, default_color=RGBColor(30, 58, 138))
                    
                    # Bottom accent border for major section headings
                    pPr = p._p.get_or_add_pPr()
                    pBdr = parse_xml(r'<w:pBdr {}><w:bottom w:val="single" w:sz="12" w:space="4" w:color="3B82F6"/></w:pBdr>'.format(nsdecls('w')))
                    pPr.append(pBdr)
            elif level == 2:
                # Subsection Heading (H2)
                p.paragraph_format.space_before = Pt(12)
                p.paragraph_format.space_after = Pt(4)
                runs = parse_inline_markdown(text)
                add_runs_to_paragraph(p, runs, default_font='Segoe UI', default_size=13.5, default_color=RGBColor(37, 99, 235))
            elif level == 3:
                # Sub-subsection Heading (H3)
                p.paragraph_format.space_before = Pt(10)
                p.paragraph_format.space_after = Pt(3)
                runs = parse_inline_markdown(text)
                add_runs_to_paragraph(p, runs, default_font='Segoe UI', default_size=12, default_color=RGBColor(51, 65, 85))
            elif level == 4:
                # Sub-sub-subsection Heading (H4)
                p.paragraph_format.space_before = Pt(8)
                p.paragraph_format.space_after = Pt(2)
                runs = parse_inline_markdown(text)
                add_runs_to_paragraph(p, runs, default_font='Calibri', default_size=11, default_color=RGBColor(71, 85, 105))
            else:
                p.paragraph_format.space_before = Pt(6)
                p.paragraph_format.space_after = Pt(2)
                runs = parse_inline_markdown(text)
                add_runs_to_paragraph(p, runs, default_font='Calibri', default_size=10.5, default_color=RGBColor(71, 85, 105))

        elif btype == 'paragraph':
            text = block['text']
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(5)
            p.paragraph_format.line_spacing = 1.15
            runs = parse_inline_markdown(text)
            add_runs_to_paragraph(p, runs, default_font='Calibri', default_size=11, default_color=RGBColor(30, 41, 59))

        elif btype == 'list_item':
            p = doc.add_paragraph()
            indent = block['indent']
            p.paragraph_format.left_indent = Inches(0.25 * indent + 0.3)
            p.paragraph_format.first_line_indent = Inches(-0.2)
            p.paragraph_format.space_before = Pt(1.5)
            p.paragraph_format.space_after = Pt(1.5)
            p.paragraph_format.line_spacing = 1.15

            # Prefix
            if block['ordered']:
                prefix_run = p.add_run(f"{block['num']}. ")
            else:
                prefix_run = p.add_run("• ")

            prefix_run.font.name = 'Calibri'
            prefix_run.font.bold = True
            prefix_run.font.size = Pt(10.5)
            prefix_run.font.color.rgb = RGBColor(37, 99, 235)

            runs = parse_inline_markdown(block['text'])
            add_runs_to_paragraph(p, runs, default_font='Calibri', default_size=11, default_color=RGBColor(30, 41, 59))

        elif btype == 'code_block':
            code_lines = block['lines']
            num_code_lines = len(code_lines)
            lang = block['lang'].upper() if block['lang'] else "CODE"

            # Optional language header tag
            header_p = doc.add_paragraph()
            header_p.paragraph_format.space_before = Pt(6)
            header_p.paragraph_format.space_after = Pt(0)
            header_p.paragraph_format.left_indent = Inches(0.25)
            header_p.paragraph_format.keep_with_next = True
            
            pPr_hdr = header_p._p.get_or_add_pPr()
            shd_hdr = parse_xml(r'<w:shd {} w:fill="E2E8F0"/>'.format(nsdecls('w')))
            pPr_hdr.append(shd_hdr)
            pBdr_hdr = parse_xml(r'<w:pBdr {}><w:left w:val="single" w:sz="18" w:space="8" w:color="2563EB"/></w:pBdr>'.format(nsdecls('w')))
            pPr_hdr.append(pBdr_hdr)
            
            hrun = header_p.add_run(f" [{lang}]")
            hrun.font.name = 'Consolas'
            hrun.font.size = Pt(8)
            hrun.font.bold = True
            hrun.font.color.rgb = RGBColor(71, 85, 105)

            for idx, line_text in enumerate(code_lines):
                p = doc.add_paragraph()
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.05
                p.paragraph_format.left_indent = Inches(0.25)

                if idx < num_code_lines - 1:
                    p.paragraph_format.keep_with_next = True
                else:
                    p.paragraph_format.space_after = Pt(6)

                pPr = p._p.get_or_add_pPr()
                # Light Gray Fill
                shd = parse_xml(r'<w:shd {} w:fill="F8FAFC"/>'.format(nsdecls('w')))
                pPr.append(shd)

                # Blue left border
                pBdr = parse_xml(r'<w:pBdr {}><w:left w:val="single" w:sz="18" w:space="8" w:color="2563EB"/></w:pBdr>'.format(nsdecls('w')))
                pPr.append(pBdr)

                run = p.add_run(line_text if line_text else " ")
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor(15, 23, 42)

        elif btype == 'table':
            headers = block['headers']
            alignments = block['alignments']
            rows = block['rows']
            num_cols = len(headers)

            if num_cols == 0:
                continue

            table = doc.add_table(rows=len(rows) + 1, cols=num_cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER

            # Table Borders XML
            tblPr = table._tbl.tblPr
            tblBorders = parse_xml(r'''
                <w:tblBorders {} >
                    <w:top w:val="single" w:sz="6" w:space="0" w:color="CBD5E1"/>
                    <w:left w:val="none"/>
                    <w:bottom w:val="single" w:sz="10" w:space="0" w:color="1E3A8A"/>
                    <w:right w:val="none"/>
                    <w:insideH w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>
                    <w:insideV w:val="none"/>
                </w:tblBorders>
            '''.format(nsdecls('w')))
            tblPr.append(tblBorders)

            # Calculate Column Widths proportionally for 6.5 in printable area
            col_max_lens = [len(h) for h in headers]
            for r in rows:
                for c_idx, cell_str in enumerate(r):
                    if c_idx < len(col_max_lens):
                        col_max_lens[c_idx] = max(col_max_lens[c_idx], len(cell_str))

            total_len = max(sum(col_max_lens), 1)
            col_widths = [Inches(max(6.5 * (l / total_len), 0.6)) for l in col_max_lens]

            # 1. Header Row
            hdr_row = table.rows[0]
            trPr = hdr_row._tr.get_or_add_trPr()
            trPr.append(parse_xml(r'<w:tblHeader {}/>'.format(nsdecls('w'))))
            trPr.append(parse_xml(r'<w:cantSplit {}/>'.format(nsdecls('w'))))

            for c_idx, h_text in enumerate(headers):
                cell = hdr_row.cells[c_idx]
                cell.width = col_widths[c_idx]
                tcPr = cell._tc.get_or_add_tcPr()
                # Navy Fill
                tcPr.append(parse_xml(r'<w:shd {} w:fill="1E3A8A"/>'.format(nsdecls('w'))))
                # Cell Padding (top/bottom 120 dxa = ~6pt, left/right 140 dxa = ~7pt)
                tcMar = parse_xml(r'''
                    <w:tcMar {}>
                        <w:top w:w="120" w:type="dxa"/>
                        <w:bottom w:w="120" w:type="dxa"/>
                        <w:left w:w="140" w:type="dxa"/>
                        <w:right w:w="140" w:type="dxa"/>
                    </w:tcMar>
                '''.format(nsdecls('w')))
                tcPr.append(tcMar)

                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
                if alignments[c_idx] == 'center':
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                elif alignments[c_idx] == 'right':
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                runs = parse_inline_markdown(h_text)
                add_runs_to_paragraph(p, runs, default_font='Segoe UI', default_size=9.5, default_color=RGBColor(255, 255, 255))
                # Ensure header runs are bold
                for r in p.runs:
                    r.font.bold = True

            # 2. Data Rows
            for r_idx, row_data in enumerate(rows, start=1):
                row = table.rows[r_idx]
                trPr = row._tr.get_or_add_trPr()
                trPr.append(parse_xml(r'<w:cantSplit {}/>'.format(nsdecls('w'))))
                row_bg = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"

                for c_idx, cell_text in enumerate(row_data):
                    cell = row.cells[c_idx]
                    cell.width = col_widths[c_idx]
                    tcPr = cell._tc.get_or_add_tcPr()
                    tcPr.append(parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), row_bg)))
                    tcMar = parse_xml(r'''
                        <w:tcMar {}>
                            <w:top w:w="100" w:type="dxa"/>
                            <w:bottom w:w="100" w:type="dxa"/>
                            <w:left w:w="140" w:type="dxa"/>
                            <w:right w:w="140" w:type="dxa"/>
                        </w:tcMar>
                    '''.format(nsdecls('w')))
                    tcPr.append(tcMar)

                    p = cell.paragraphs[0]
                    p.paragraph_format.space_before = Pt(0)
                    p.paragraph_format.space_after = Pt(0)
                    p.paragraph_format.line_spacing = 1.1

                    if alignments[c_idx] == 'center':
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    elif alignments[c_idx] == 'right':
                        p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                    else:
                        p.alignment = WD_ALIGN_PARAGRAPH.LEFT

                    runs = parse_inline_markdown(cell_text)
                    add_runs_to_paragraph(p, runs, default_font='Calibri', default_size=9.5, default_color=RGBColor(30, 41, 59))

            # Add spacing paragraph after table
            sp = doc.add_paragraph()
            sp.paragraph_format.space_before = Pt(0)
            sp.paragraph_format.space_after = Pt(6)

        elif btype == 'hr':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(10)
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(r'<w:pBdr {}><w:bottom w:val="single" w:sz="6" w:space="1" w:color="CBD5E1"/></w:pBdr>'.format(nsdecls('w')))
            pPr.append(pBdr)

    doc.save(output_filepath)
    print(f"Document successfully created at: {output_filepath}")


def main():
    input_file = r"d:\Hospital MYSQL Databse\project_report.md"
    output_file = r"d:\Hospital MYSQL Databse\Hospital_Management_System_Report.docx"

    print(f"Reading markdown input from: {input_file}")
    with open(input_file, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    print("Parsing markdown blocks...")
    blocks = parse_markdown_blocks(lines)
    print(f"Parsed {len(blocks)} block elements.")

    print(f"Generating DOCX document: {output_file}")
    build_docx(blocks, output_file)
    print("Done!")

if __name__ == '__main__':
    main()
