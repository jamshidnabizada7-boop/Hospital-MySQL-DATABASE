import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import nsdecls, qn

doc = docx.Document()

# Set margins
sections = doc.sections
for section in sections:
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

# Header & Footer
header = doc.sections[0].header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
hrun = hp.add_run("Hospital Management System | Technical Report")
hrun.font.name = 'Calibri'
hrun.font.size = Pt(8.5)
hrun.font.color.rgb = RGBColor(100, 116, 139)

footer = doc.sections[0].footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.LEFT
frun = fp.add_run("Enterprise HMS Project Report")
frun.font.name = 'Calibri'
frun.font.size = Pt(8.5)
frun.font.color.rgb = RGBColor(100, 116, 139)

# Test code block formatting
p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(4)
p.paragraph_format.space_after = Pt(4)
p.paragraph_format.left_indent = Inches(0.25)
pPr = p._p.get_or_add_pPr()

# Shading
shd = parse_xml(r'<w:shd {} w:fill="F1F5F9"/>'.format(nsdecls('w')))
pPr.append(shd)

# Left border
pBdr = parse_xml(r'<w:pBdr {}><w:left w:val="single" w:sz="18" w:space="8" w:color="2563EB"/></w:pBdr>'.format(nsdecls('w')))
pPr.append(pBdr)

run = p.add_run("SELECT * FROM Patient WHERE Is_Active = 1;")
run.font.name = 'Consolas'
run.font.size = Pt(9.5)
run.font.color.rgb = RGBColor(15, 23, 42)

# Test Table
table = doc.add_table(rows=3, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER

# Set Table Borders XML
tblPr = table._tbl.tblPr
tblBorders = parse_xml(r'''
    <w:tblBorders {} >
        <w:top w:val="single" w:sz="6" w:space="0" w:color="CBD5E1"/>
        <w:left w:val="none"/>
        <w:bottom w:val="single" w:sz="8" w:space="0" w:color="94A3B8"/>
        <w:right w:val="none"/>
        <w:insideH w:val="single" w:sz="4" w:space="0" w:color="E2E8F0"/>
        <w:insideV w:val="none"/>
    </w:tblBorders>
'''.format(nsdecls('w')))
tblPr.append(tblBorders)

# Header Row
hdr_row = table.rows[0]
trPr = hdr_row._tr.get_or_add_trPr()
trPr.append(parse_xml(r'<w:tblHeader {}/>'.format(nsdecls('w'))))
trPr.append(parse_xml(r'<w:cantSplit {}/>'.format(nsdecls('w'))))

headers = ["Col 1", "Col 2", "Col 3"]
for i, cell in enumerate(hdr_row.cells):
    cell.text = headers[i]
    tcPr = cell._tc.get_or_add_tcPr()
    tcPr.append(parse_xml(r'<w:shd {} w:fill="1E3A8A"/>'.format(nsdecls('w'))))
    # Cell margins (top 120 dxa, bottom 120 dxa, left 150 dxa, right 150 dxa)
    tcMar = parse_xml(r'''
        <w:tcMar {}>
            <w:top w:w="120" w:type="dxa"/>
            <w:bottom w:w="120" w:type="dxa"/>
            <w:left w:w="150" w:type="dxa"/>
            <w:right w:w="150" w:type="dxa"/>
        </w:tcMar>
    '''.format(nsdecls('w')))
    tcPr.append(tcMar)
    # Format text in cell
    for p in cell.paragraphs:
        for r in p.runs:
            r.font.name = 'Calibri'
            r.font.bold = True
            r.font.color.rgb = RGBColor(255, 255, 255)
            r.font.size = Pt(10)

# Data Rows
for r_idx, row in enumerate(table.rows[1:], start=1):
    trPr = row._tr.get_or_add_trPr()
    trPr.append(parse_xml(r'<w:cantSplit {}/>'.format(nsdecls('w'))))
    fill_color = "F8FAFC" if r_idx % 2 == 1 else "FFFFFF"
    for c_idx, cell in enumerate(row.cells):
        cell.text = f"Data {r_idx},{c_idx}"
        tcPr = cell._tc.get_or_add_tcPr()
        tcPr.append(parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), fill_color)))
        tcMar = parse_xml(r'''
            <w:tcMar {}>
                <w:top w:w="100" w:type="dxa"/>
                <w:bottom w:w="100" w:type="dxa"/>
                <w:left w:w="150" w:type="dxa"/>
                <w:right w:w="150" w:type="dxa"/>
            </w:tcMar>
        '''.format(nsdecls('w')))
        tcPr.append(tcMar)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.name = 'Calibri'
                r.font.size = Pt(9.5)
                r.font.color.rgb = RGBColor(30, 41, 59)

doc.save("test_advanced.docx")
print("Saved test_advanced.docx successfully!")
