import os
import sys
import zipfile
import xml.etree.ElementTree as ET

def verify_docx(file_path):
    print(f"=== VERIFYING DOCX REPORT: {file_path} ===")
    results = {}

    # Check 1: Existence and File Size
    if not os.path.exists(file_path):
        print(f"FAIL: File does not exist: {file_path}")
        return False, results
    
    file_size = os.path.getsize(file_path)
    print(f"PASS: File exists. Size: {file_size} bytes ({file_size / 1024:.2f} KB)")
    results['file_exists'] = True
    results['file_size_bytes'] = file_size
    
    if file_size <= 0:
        print("FAIL: File size is 0 bytes.")
        return False, results

    # Check 2: Valid Zip Archive
    if not zipfile.is_zipfile(file_path):
        print("FAIL: File is not a valid zip archive (.docx format requirement).")
        return False, results
    
    print("PASS: File is a valid zip archive.")
    results['is_zip'] = True

    try:
        with zipfile.ZipFile(file_path, 'r') as z:
            namelist = z.namelist()
            print(f"Zip entries count: {len(namelist)}")
            if 'word/document.xml' not in namelist:
                print("FAIL: 'word/document.xml' not found in zip archive.")
                return False, results
            
            doc_xml = z.read('word/document.xml')
            media_files = [f for f in namelist if f.startswith('word/media/')]
            results['media_files'] = media_files
            print(f"Media files count: {len(media_files)}")
    except Exception as e:
        print(f"FAIL: Error reading zip archive: {e}")
        return False, results

    # Check 3 & 4: Parsing with python-docx if available, else ElementTree
    has_python_docx = False
    paragraphs_list = []
    tables_list = []
    runs_count = 0
    full_text = ""

    try:
        import docx
        has_python_docx = True
        print("python-docx package is available. Parsing using python-docx...")
        doc = docx.Document(file_path)
        
        for p in doc.paragraphs:
            paragraphs_list.append(p.text)
            runs_count += len(p.runs)
            full_text += p.text + "\n"
            
        for t in doc.tables:
            table_data = []
            for row in t.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            tables_list.append(table_data)

    except ImportError:
        print("python-docx not installed. Parsing XML directly with ElementTree...")
        root = ET.fromstring(doc_xml)
        namespaces = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
        
        for p in root.findall('.//w:p', namespaces):
            p_text = "".join([t.text for t in p.findall('.//w:t', namespaces) if t.text])
            paragraphs_list.append(p_text)
            runs_count += len(p.findall('.//w:r', namespaces))
            full_text += p_text + "\n"
            
        for tbl in root.findall('.//w:tbl', namespaces):
            table_data = []
            for tr in tbl.findall('.//w:tr', namespaces):
                row_data = []
                for tc in tr.findall('.//w:tc', namespaces):
                    cell_text = "".join([t.text for t in tc.findall('.//w:t', namespaces) if t.text])
                    row_data.append(cell_text.strip())
                table_data.append(row_data)
            tables_list.append(table_data)

    except Exception as e:
        print(f"FAIL: Error parsing document structure: {e}")
        return False, results

    results['paragraphs_count'] = len(paragraphs_list)
    results['tables_count'] = len(tables_list)
    results['runs_count'] = runs_count
    results['word_count'] = len(full_text.split())

    print(f"PASS: Parsed document successfully.")
    print(f" - Paragraphs count: {len(paragraphs_list)}")
    print(f" - Tables count: {len(tables_list)}")
    print(f" - Runs count: {runs_count}")
    print(f" - Total word count: {results['word_count']}")

    # Check 5: Section Headings Verification
    required_headings = [
        "Abstract",
        "System Architecture",
        "Database Schema",
        "Access Control",
        "Frontend UI Flow",
        "Future Enhancements"
    ]

    print("\n--- Checking Required Section Headings ---")
    missing_headings = []
    found_headings = {}

    for req in required_headings:
        match_found = False
        # Search paragraphs for exact or phrase matches
        for p in paragraphs_list:
            if req.lower() in p.lower():
                match_found = True
                found_headings[req] = p
                break
        if match_found:
            sample_text = found_headings[req][:80].replace('\n', ' ')
            print(f"[OK] Found requirement '{req}': matching text: '{sample_text}...'")
        else:
            print(f"[MISSING] MISSING requirement '{req}'")
            missing_headings.append(req)

    results['found_headings'] = found_headings
    results['missing_headings'] = missing_headings

    if missing_headings:
        print(f"\nFAIL: Missing required section headings: {missing_headings}")
        return False, results

    print("\n=== ALL EMPIRICAL VERIFICATIONS PASSED ===")
    return True, results

if __name__ == "__main__":
    docx_path = r"d:\Hospital MYSQL Databse\Hospital_Management_System_Report.docx"
    success, details = verify_docx(docx_path)
    print("\nDetailed Summary:")
    print(details)
    if success:
        sys.exit(0)
    else:
        sys.exit(1)
