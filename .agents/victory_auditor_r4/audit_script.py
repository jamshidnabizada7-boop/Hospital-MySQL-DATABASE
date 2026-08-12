import os
import re
import sys
import docx

def run_victory_audit():
    root_dir = r"d:\Hospital MYSQL Databse"
    md_path = os.path.join(root_dir, "project_report.md")
    docx_path = os.path.join(root_dir, "Hospital_Management_System_Report.docx")
    
    report = []
    report.append("==================================================")
    report.append("  INDEPENDENT VICTORY AUDIT PROGRAMMATIC ANALYSIS ")
    report.append("==================================================\n")
    
    # ---------------------------------------------------------
    # PHASE 1: TIMELINE & ARTIFACT VERIFICATION
    # ---------------------------------------------------------
    report.append("--- PHASE 1: Timeline & Artifact Verification ---")
    md_exists = os.path.isfile(md_path)
    docx_exists = os.path.isfile(docx_path)
    
    md_size = os.path.getsize(md_path) if md_exists else 0
    docx_size = os.path.getsize(docx_path) if docx_exists else 0
    
    report.append(f"Markdown file exists: {md_exists} (Path: {md_path}, Size: {md_size:,} bytes)")
    report.append(f"DOCX file exists: {docx_exists} (Path: {docx_path}, Size: {docx_size:,} bytes)")
    
    p1_pass = md_exists and docx_exists and md_size > 10000 and docx_size > 10000
    report.append(f"Phase 1 Artifact Check: {'PASS' if p1_pass else 'FAIL'}\n")
    
    # ---------------------------------------------------------
    # PHASE 2: ANTI-CHEATING & QUALITY AUDIT
    # ---------------------------------------------------------
    report.append("--- PHASE 2: Anti-Cheating & Quality Audit ---")
    
    with open(md_path, "r", encoding="utf-8") as f:
        md_text = f.read()
        
    word_count = len(re.findall(r'\w+', md_text))
    char_count = len(md_text)
    line_count = len(md_text.splitlines())
    
    report.append(f"Markdown Statistics: {line_count:,} lines, {word_count:,} words, {char_count:,} characters")
    
    # Placeholder checks
    placeholders = [
        r'\bTODO\b', r'\bFIXME\b', r'\[Insert', r'\[insert', r'\bTBD\b', 
        r'Lorem Ipsum', r'placeholder', r'sample text', r'<insert'
    ]
    found_placeholders = []
    for ph in placeholders:
        matches = re.findall(ph, md_text, re.IGNORECASE)
        if matches:
            found_placeholders.append((ph, len(matches)))
            
    report.append(f"Placeholder scan results: {len(found_placeholders)} forbidden patterns found")
    for ph, cnt in found_placeholders:
        report.append(f"  - Pattern '{ph}': {cnt} occurrences")
        
    # Technical depth checks (Database focus)
    sql_blocks = re.findall(r'```sql(.*?)```', md_text, re.DOTALL | re.IGNORECASE)
    all_code_blocks = re.findall(r'```(.*?)```', md_text, re.DOTALL)
    
    sql_keywords = ['CREATE TABLE', 'FOREIGN KEY', 'PRIMARY KEY', 'JOIN', 'GROUP BY', 'TRANSACTION', 'INDEX', 'TRIGGER', 'VIEW', 'SELECT']
    sql_kw_counts = {}
    for kw in sql_keywords:
        sql_kw_counts[kw] = len(re.findall(re.escape(kw), md_text, re.IGNORECASE))
        
    report.append(f"Total Code Blocks: {len(all_code_blocks)}")
    report.append(f"SQL Code Blocks: {len(sql_blocks)}")
    report.append("Database Keyword Frequency:")
    for kw, cnt in sql_kw_counts.items():
        report.append(f"  - {kw}: {cnt}")
        
    # Required sections check
    required_sections = [
        ("Abstract", [r'abstract', r'executive summary']),
        ("System Architecture", [r'system architecture']),
        ("Database Schema", [r'database schema', r'entity relationship', r'er diagram', r'er structure']),
        ("Access Control", [r'access control', r'security features', r'rbac']),
        ("Frontend UI Flow", [r'frontend ui flow', r'frontend architecture', r'ui flow']),
        ("Future Enhancements", [r'future enhancements', r'conclusion'])
    ]
    
    missing_sections = []
    for sec_name, patterns in required_sections:
        found = any(re.search(pat, md_text, re.IGNORECASE) for pat in patterns)
        if not found:
            missing_sections.append(sec_name)
            
    report.append(f"Missing required sections in MD: {missing_sections if missing_sections else 'None'}")
    
    p2_pass = (len(found_placeholders) == 0) and (word_count > 3000) and (len(sql_blocks) >= 3) and (len(missing_sections) == 0)
    report.append(f"Phase 2 Anti-Cheating & Quality Check: {'PASS' if p2_pass else 'FAIL'}\n")
    
    # ---------------------------------------------------------
    # PHASE 3: EMPIRICAL TEST & FILE EXECUTION AUDIT
    # ---------------------------------------------------------
    report.append("--- PHASE 3: Empirical Test & File Execution Audit ---")
    
    try:
        doc = docx.Document(docx_path)
        p_count = len(doc.paragraphs)
        t_count = len(doc.tables)
        
        docx_text = []
        for p in doc.paragraphs:
            if p.text.strip():
                docx_text.append(p.text.strip())
                
        table_cell_count = 0
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    table_cell_count += 1
                    if cell.text.strip():
                        docx_text.append(cell.text.strip())
                        
        full_docx_str = "\n".join(docx_text)
        docx_word_count = len(re.findall(r'\w+', full_docx_str))
        
        # Check XML styling elements (borders for headings/code blocks)
        h1_border_count = 0
        code_shaded_count = 0
        for p in doc.paragraphs:
            pPr = p._p.get_or_add_pPr()
            if pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}pBdr') is not None:
                h1_border_count += 1
            if pPr.find('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd') is not None:
                code_shaded_count += 1
                
        report.append(f"DOCX loaded successfully!")
        report.append(f"Paragraphs count: {p_count:,}")
        report.append(f"Tables count: {t_count} (Total cells: {table_cell_count})")
        report.append(f"Total extracted words from DOCX: {docx_word_count:,}")
        report.append(f"XML Borders / Accents applied: {h1_border_count:,}")
        report.append(f"XML Shaded Code Block Lines: {code_shaded_count:,}")
        
        # Check required sections in DOCX
        missing_docx_sections = []
        for sec_name, patterns in required_sections:
            found = any(re.search(pat, full_docx_str, re.IGNORECASE) for pat in patterns)
            if not found:
                missing_docx_sections.append(sec_name)
                
        report.append(f"Missing required sections in DOCX: {missing_docx_sections if missing_docx_sections else 'None'}")
        
        p3_pass = (p_count > 100) and (t_count >= 2) and (docx_word_count > 3000) and (len(missing_docx_sections) == 0)
        report.append(f"Phase 3 File Execution Check: {'PASS' if p3_pass else 'FAIL'}\n")
        
    except Exception as e:
        report.append(f"ERROR loading/reading DOCX file: {e}")
        p3_pass = False
        report.append("Phase 3 File Execution Check: FAIL\n")
        
    # OVERALL VERDICT
    overall_pass = p1_pass and p2_pass and p3_pass
    verdict = "VICTORY CONFIRMED" if overall_pass else "VICTORY REJECTED"
    
    report.append("==================================================")
    report.append(f"OVERALL AUDIT VERDICT: {verdict}")
    report.append("==================================================")
    
    out_str = "\n".join(report)
    print(out_str)
    
    summary_path = os.path.join(root_dir, r".agents\victory_auditor_r4\audit_summary.txt")
    with open(summary_path, "w", encoding="utf-8") as f:
        f.write(out_str)

if __name__ == "__main__":
    run_victory_audit()
